from output import *
import re
import docx2txt
from docx import Document
import pandas as pd
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm,Pt
from docx.oxml.ns import qn
from operator import itemgetter
from image import *
import os
import shutil
import zipfile
import fnmatch

def replace_text(filename,old,new):
    doc = Document(filename)
    for p in doc.paragraphs:
        if old in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if old in inline[i].text:
                    text = inline[i].text.replace(old, new)
                    inline[i].text = text
    doc.save('test.docx')
    return
def replace_word():
    if park.loc['是否用氢'][1]=='是':
        #不离网
        if str(grid_planning_output_json["flag_isloate"])=="0":
            if park.loc['园区规范'][1]=='无':
                replace_text('./输出文档-无离网.docx', '园区规范','')
            else:
                replace_text('./输出文档-无离网.docx', '园区规范','（3）'+park.loc['园区名称'][1]+'管委会提供的基础资料：'+park.loc['园区规范'][1])
        else:
            if park.loc['园区规范'][1]=='无':
                replace_text('./输出文档.docx', '园区规范','')
            else:
                replace_text('./输出文档.docx', '园区规范','（3）'+park.loc['园区名称'][1]+'管委会提供的基础资料：'+park.loc['园区规范'][1])
    else:
        if park.loc['园区规范'][1]=='无':
            replace_text('./输出文档-无氢.docx', '园区规范','')
        else:
            replace_text('./输出文档-无氢.docx', '园区规范','（3）'+park.loc['园区名称'][1]+'管委会提供的基础资料：'+park.loc['园区规范'][1])

    replace_text('test.docx','（园区名称）',park.loc['园区名称'][1])

    replace_text('test.docx','description',park.loc['园区描述'][1])

    replace_text('test.docx','城市描述',city.loc['城市描述'][1])


    replace_text('test.docx','土地使用情况',park.loc['土地使用情况'][1])

    replace_text('test.docx','位置描述',park.loc['位置描述'][1])

    replace_text('test.docx','城市名称',city.loc['城市名称'][1])

    replace_text('test.docx','平均温度',str(city.loc['平均温度'][1]))

    replace_text('test.docx','年最高温度',str(city.loc['年最高温度'][1]))

    replace_text('test.docx','年最低温度',str(city.loc['年最低温度'][1]))

    replace_text('test.docx','气候描述',city.loc['气候描述'][1])

    replace_text('test.docx','气候分区',city.loc['气候分区'][1])

    # replace_text('test.docx','光伏情况',city.loc['光伏情况'][1])

    replace_text('test.docx','采暖供冷描述',city.loc['采暖供冷期描述'][1])

    replace_text('test.docx','电价描述',park.loc['电价描述'][1])

    replace_text('test.docx','用能政策',park.loc['用能政策'][1])
    if park.loc['地热资源评价'][1]=='无':
        replace_text('test.docx', '地热资源评价','')
    else:
        replace_text('test.docx','地热资源评价','地热资源属新能源，价廉、方便且无污染，可广泛应用于化工、纺织工业、居民区供热及温室栽培。高温热水中的氟、硅酸、碘、硼酸等含量均达到或超过医疗矿水含量，对多种疾病具有良好的理疗效果，有较高医疗价值。根据中国建科院《中国地源热泵应用适宜性评价》结果显示：寒冷气候区为适宜区，表明各项指标均适宜；夏热冬冷气候区为一般适宜区，表明吸排热量不平衡率偏高、且与当地常规系统相比经济性较差。'+park.loc['地热资源评价'][1])

    replace_text('test.docx','permits',park.loc['卖电许可'][1])

    replace_text('test.docx','氢价',str(park.loc['氢价'][1]))

    if park.loc['卖电许可'][1]=='不允许':
        replace_text('test.docx','，以及卖电收益产生的抵扣。','。')

    replace_text('test.docx','供能对象描述',park.loc['供能对象描述'][1])

    replace_text('test.docx','供能方案描述',park.loc['用能方案描述'][1])

    replace_text('test.docx','所在省份',city.loc['所在省份'][1])

    replace_text('test.docx','供能方案描述',park.loc['用能方案描述'][1])

    replace_text('test.docx','制氢潜力',park.loc['制氢潜力'][1])

    replace_text('test.docx',"aaa",grid_planning_output_json['equipment_cost'])
    replace_text('test.docx','bbb',grid_operation_output_json['cer'])
    replace_text('test.docx','ccc',grid_operation_output_json['cer_perm2'])

    replace_text('test.docx','龘',itgrid_planning_output_json['equipment_cost'])
    replace_text('test.docx','薅',itgrid_operation_output_json['cer'])
    replace_text('test.docx','ic',itgrid_operation_output_json['cer_perm2'])

    replace_text('test.docx','eee',grid_planning_output_json['receive_year'])
    replace_text('test.docx','ite',itgrid_planning_output_json['receive_year'])

    replace_text('test.docx','fff',grid_operation_output_json['cer_rate'])
    replace_text('test.docx','if',itgrid_operation_output_json['cer_rate'])


def change_table_air_para(doc):
    tables = doc.tables[0]
    for i in range(len(tables.rows)):
        # 在第该表格i行列的单元格内输入对应内容”
        run = tables.cell(i, 1).paragraphs[0].add_run(str(list(gongnuan[1])[i]))
        # 设置字体
        run.font.name = 'Times New Roman'
        # 字体大小
        run.font.size = Pt(12)
        r = run._element
        # 中文宋体
        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 居中
        tables.cell(i, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def change_table_equip_para(doc):
    main_input=main_input_json["device"]
    tables = doc.tables[1]
    para=[]
    #氢压机
    if park.loc["是否用氢"][1] == "否" or main_input["co"]["power_max"]==0:
        para+=["","",""]
    else:
        para+=[main_input["co"]["beta_co"],
               main_input["co"]["cost"],
               main_input["co"]["crf"]]
    #燃料电池
    if main_input["fc"]["power_max"]==0:
        para+=["","","","",""]
    else:
        para+=[main_input["fc"]["eta_fc_p"],
                main_input["fc"]["eta_ex_g"],
                main_input["fc"]["theta_ex"],
                main_input["fc"]["cost"],
                main_input["fc"]["crf"]]
    #热水罐
    if main_input["ht"]["water_max"]==0:
        para+=["","","",""]
    else:
        para+=[main_input["ht"]["t_max"],
               main_input["ht"]["t_min"],
               main_input["ht"]["cost"],
               main_input["ht"]["crf"]]
    #电锅炉
    if main_input["eb"]["power_max"]==0:
        para+=["","",""]
    else:
        para+=[main_input["eb"]["beta_eb"],
               main_input["eb"]["cost"],
               main_input["eb"]["crf"]]
    #空气源热泵
    if main_input["hp"]["power_max"]==0:
        para+=["","","",""]
    else:
        para+=[main_input["hp"]["beta_hpg"],
               main_input["hp"]["beta_hpq"],
               main_input["hp"]["cost"],
               main_input["hp"]["crf"]]
    #地源热泵
    if main_input["ghp"]["power_max"]==0:
        para+=["","","",""]
    else:
        para+=[main_input["ghp"]["beta_ghpg"],
               main_input["ghp"]["beta_ghpq"],
               main_input["ghp"]["cost"],
               main_input["ghp"]["crf"]]
    #浅层地热井
    if main_input["gtw"]["number_max"]==0:
        para+=["","",""]
    else:
        para+=[main_input["gtw"]["number_max"],
               main_input["gtw"]["cost"],
               main_input["gtw"]["crf"]]
    #冷水罐
    if main_input["ct"]["water_max"]==0:
        para+=["","","",""]
    else:
        para+=[main_input["ct"]["t_max"],
               main_input["ct"]["t_min"],
               main_input["ct"]["cost"],
               main_input["ct"]["crf"]]
    #储氢罐
    if park.loc["是否用氢"][1] == "否" or main_input["hst"]["sto_max"] == 0:
        para += ["", "", ""]
    else:
        para+=[main_input["hst"]["cost"],
               main_input["hst"]["crf"],
               main_input["hst"]["sto_max"]]
    #电解槽
    if main_input["el"]["power_max"]==0 or main_input["el"]["nm3_max"]==0 :
        para+=["","",""]
    else:
        para+=[main_input["el"]["cost"]*50/11.2,
               main_input["el"]["crf"],
               main_input["el"]["power_max"]]
    #光伏板
    if main_input["pv"]["area_max"]==0:
        para+=["",""]
    else:
        para+=[main_input["pv"]["cost"],
               main_input["pv"]["crf"]]
    #太阳能集热器
    if main_input["sc"]["area_max"]==0:
        para+=["",""]
    else:
        para+=[main_input["sc"]["cost"],
               main_input["sc"]["crf"]]
    s=[str(i) for i in para]
    for i in range(1, len(tables.rows)):
        if s[i - 1] != "":
            # 在第该表格i行2列的单元格内输入对应内容”
            run = tables.cell(i, 2).paragraphs[0].add_run(s[i - 1])
            # 设置字体
            run.font.name = 'Times New Roman'
            # 字体大小
            run.font.size = Pt(12)
            r = run._element
            # 中文宋体
            r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 居中
            tables.cell(i, 2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #删除空白行
    for row in tables.rows:
        if not row.cells[2].text:
            row._element.getparent().remove(row._element)


def change_table_equip_allocation(tables, json):
    keys = ["p_fc_max",
            "num_gtw",
            "p_hpg_max",
            "p_hp_max",
            "p_eb_max",
            "p_el_max",
            "hst",
            "m_ht",
            "m_ct",
            "area_pv",
            "area_sc",
            "p_co"]
    replace = [int(float(i)) for i in list(itemgetter(*keys)(json))]

    #储热罐、储冷罐kg改成t
    if replace[7]>1000:
        replace[7]=replace[7]//1000
        p = tables.cell(8, 3).paragraphs[0]
        inline = p.runs
        # Loop added to work with runs (strings with same style)
        for i in range(len(inline)):
            text = inline[i].text.replace('kg', 't')
            inline[i].text = text
    if replace[8]>1000:
        replace[8]=replace[8]//1000
        p = tables.cell(9, 3).paragraphs[0]
        inline = p.runs
        # Loop added to work with runs (strings with same style)
        for i in range(len(inline)):
            text = inline[i].text.replace('kg', 't')
            inline[i].text = text

    for i in range(1,len(tables.rows)):
        if replace[i-1] !=0:
            # 在第该表格i行列的单元格内输入对应内容”
            run = tables.cell(i, 2).paragraphs[0].add_run(str(replace[i-1]))
            # 设置字体
            run.font.name = 'Times New Roman'
            # 字体大小
            run.font.size = Pt(12)
            r = run._element
            # 中文宋体
            r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 居中
            tables.cell(i, 2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def change_table_eco_analyse(tables, planing_json,operation_json):
    replace = [int(float(planing_json['equipment_cost'])),
               int(abs(float(operation_json['operation_cost']))),
               float(operation_json['cost_save_rate']),
               float(planing_json['receive_year']),
               int(float(operation_json['co2'])),
               float(operation_json['cer_rate']),
               float(operation_json['revenue'])]
    for i in range(1,len(tables.rows)):
        run = tables.cell(i, 1).paragraphs[0].add_run(str(replace[i - 1]))
        # 设置字体
        run.font.name = 'Times New Roman'
        # 字体大小
        run.font.size = Pt(12)
        r = run._element
        # 中文宋体
        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 居中
        tables.cell(i, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #收益/成本
    if float(operation_json['operation_cost'])<0:
        p=tables.cell(2,0).paragraphs[0]
        inline = p.runs
        # Loop added to work with runs (strings with same style)
        for i in range(len(inline)):
                text = inline[i].text.replace('年化运行成本', '年化运行收益')
                inline[i].text = text

def zip_docx(startdir, file_news):
    z = zipfile.ZipFile(file_news, 'w', zipfile.ZIP_DEFLATED)
    for dirpath, dirnames, filenames in os.walk(startdir):
        fpath = dirpath.replace(startdir, '')
        fpath = fpath and fpath + os.sep or ''
        for filename in filenames:
            z.write(os.path.join(dirpath, filename), fpath+filename)
def dynamicTable(table):
    #删除空白行
    for row in table.rows:
        if not row.cells[2].text:
            row._element.getparent().remove(row._element)
    #重编序号
    for i in range(1,len(table.rows)):
        run = table.cell(i, 0).paragraphs[0].add_run(str(i))
        # 设置字体
        run.font.name = 'Times New Roman'
        # 字体大小
        run.font.size = Pt(12)
        r = run._element
        # 中文宋体
        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 居中
        table.cell(i, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def replace_equip_text(table,replace_str):
    s = ''
    for i in range(1, len(table.rows)):
        s = s + table.cell(i, 1).text
    #替换文中对应文本
    replace=s.replace('数目','、').replace('容量','、').replace('功率','、').replace('面积','、').strip('、')
    replace_text('test.docx',replace_str,replace)

if __name__=='__main__':
    ###读配置文件
    park = pd.read_excel('./配置文档数据.xlsx', sheet_name='园区文字', header=None, index_col=0)
    city = pd.read_excel('./配置文档数据.xlsx', sheet_name='城市文字', header=None, index_col=0)
    gongnuan = pd.read_excel('./配置文档数据.xlsx', sheet_name='供暖规范', header=None, index_col=0)

    ###更改文字与表格
    replace_word()
    doc = Document('test.docx')

    change_table_air_para(doc)
    change_table_equip_para(doc)

    table2=doc.tables[2]
    table3 = doc.tables[3]
    if park.loc['是否用氢'][1] == '是' and str(grid_planning_output_json["flag_isloate"])!="0":
        table4 = doc.tables[4]
        table5 = doc.tables[5]

    change_table_equip_allocation(table2,grid_planning_output_json)
    if park.loc['是否用氢'][1] == '是' and str(grid_planning_output_json["flag_isloate"])!="0":
        change_table_equip_allocation(table4, itgrid_planning_output_json)
    change_table_eco_analyse(table3, grid_planning_output_json, grid_operation_output_json)
    if park.loc['是否用氢'][1] == '是' and str(grid_planning_output_json["flag_isloate"])!="0":
        change_table_eco_analyse(table5, itgrid_planning_output_json, itgrid_operation_output_json)
    dynamicTable(table2)
    if park.loc['是否用氢'][1] == '是' and str(grid_planning_output_json["flag_isloate"])!="0":
        dynamicTable(table4)
    doc.save('test.docx')

    #更改对应设备文字
    replace_equip_text(table2,'grid-equipment')
    if park.loc['是否用氢'][1] == '是' and str(grid_planning_output_json["flag_isloate"])!="0":
        replace_equip_text(table4, 'off-equipment')


    ###更改图片,判断电热冷

    ##只有电热
    if sum(grid_planning_output_json['q_demand']) == 0:
        replace_text('test.docx', 'load condition',
                     '全年的电-热需求呈规律性变化趋势，其中：电负荷峰值为{:.0f}kWh，热负荷峰值为{:.0f}kWh。全年电负荷共{:.0f}kWh，热负荷共{:.0f}kWh。'.format(
                         grid_planning_output_json['ele_load_max'], grid_planning_output_json['g_demand_max'],
                          grid_planning_output_json['ele_load_sum'],
                         grid_planning_output_json['g_demand_sum']))
    ##只有电冷
    if sum(grid_planning_output_json['g_demand']) == 0:
        replace_text('test.docx', 'load condition',
                     '全年的电-冷需求呈规律性变化趋势，其中：电负荷峰值为{:.0f}kWh，冷负荷峰值为{:.0f}kWh。全年电负荷共{:.0f}kWh，冷负荷共{:.0f}kWh。'.format(
                         grid_planning_output_json['ele_load_max'], grid_planning_output_json['q_demand_max'],
                          grid_planning_output_json['ele_load_sum'],
                         grid_planning_output_json['q_demand_sum']))
    ##电热冷都有
    else:
        replace_text('test.docx','load condition',
                     '全年的电-热-冷需求呈规律性变化趋势，其中：电负荷峰值为{:.0f}kWh，热负荷峰值为{:.0f}kWh，冷负荷峰值为{:.0f}kWh。全年电负荷共{:.0f}kWh，热负荷共{:.0f}kWh，冷负荷共{:.0f}kWh。'.format(
                         grid_planning_output_json['ele_load_max'],grid_planning_output_json['g_demand_max'],
                         grid_planning_output_json['q_demand_max'],grid_planning_output_json['ele_load_sum'],
                         grid_planning_output_json['g_demand_sum'],grid_planning_output_json['q_demand_sum']))


    with zipfile.ZipFile('./test.docx') as z:
        z.extractall('./docx/')  # 解压docx文件
    light_image()
    load_image()
    load_image_1_01()
    load_image_7_15()
    # ##只有电冷
    # if sum(grid_planning_output_json['g_demand']) == 0:
    #     os.remove('.\docx\word\media\image4.png')
    # ##只有电热
    # if sum(grid_planning_output_json['q_demand']) == 0:
    #     os.remove('.\docx\word\media\image5.png')

    zip_docx('./docx/', 'out.docx')

    ###删除过程文件
    os.remove('./test.docx')
    shutil.rmtree('./docx')



